import io

from docx import Document

from api import exam_generator as exam_api


def test_generate_exam_word_uses_ministry_header_and_forces_multiple_choice(client_factory, db_session, seed):
    teacher = seed.user(db_session, "teacher.exam@example.com", role="teacher", full_name="Exam Teacher")
    subject = seed.subject(db_session, "Cơ sở Hệ điều hành")
    classroom = seed.classroom(db_session, "Hệ điều hành 01", subject, teacher, class_code="CLS-OS01")
    seed.document(
        db_session,
        subject,
        classroom,
        teacher,
        filename="os.pdf",
        title="Giáo trình Hệ điều hành",
        visible=True,
    )
    seed.chunk(
        db_session,
        subject,
        "os.pdf",
        "Hệ điều hành quản lý tài nguyên phần cứng, tiến trình, bộ nhớ và cung cấp dịch vụ cho chương trình ứng dụng.",
        classroom=classroom,
    )
    seed.question(
        db_session,
        subject,
        content="Chức năng chính của hệ điều hành là gì?",
        options=[
            "A. Quản lý tài nguyên và cung cấp dịch vụ cho chương trình",
            "B. Chỉ dùng để soạn thảo văn bản",
            "C. Chỉ dùng để lướt web",
            "D. Chỉ dùng để chơi game",
        ],
        correct_answer="A",
        source_file="os.pdf",
        explanation="Hệ điều hành quản lý tài nguyên và làm trung gian giữa phần cứng với ứng dụng.",
    )

    client = client_factory((exam_api.router, "/api/exam"))

    response = client.post(
        "/api/exam/generate-word",
        json={
            "class_id": classroom.id,
            "subject": subject.name,
            "exam_type": "tự luận",
            "num_questions": 1,
            "num_versions": 1,
            "level": "Trung bình",
        },
    )

    assert response.status_code == 200
    assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    generated_doc = Document(io.BytesIO(response.content))
    header_table = generated_doc.tables[0]
    header_left = header_table.cell(0, 0).text
    header_right = header_table.cell(0, 1).text
    paragraphs = [paragraph.text for paragraph in generated_doc.paragraphs if paragraph.text.strip()]

    assert "BỘ GIÁO DỤC VÀ ĐÀO TẠO" in header_left
    assert "TRƯỜNG ĐẠI HỌC XÂY DỰNG HÀ NỘI" in header_left
    assert "KHOA CÔNG NGHỆ THÔNG TIN" in header_left
    assert "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM" in header_right
    assert "Độc lập - Tự do - Hạnh phúc" in header_right
    assert "ĐỀ THI KẾT THÚC HỌC PHẦN" in paragraphs

    info_paragraph = next(text for text in paragraphs if "Môn thi:" in text)
    assert "Môn thi: Cơ sở Hệ điều hành" in info_paragraph
    assert "Hình thức thi: Trắc Nghiệm" in info_paragraph
    assert "Mã đề thi:" in info_paragraph
    assert "Họ và tên sinh viên:" in info_paragraph

    assert any(text.startswith("Câu 1:") for text in paragraphs)
    assert any(text.startswith("A. ") for text in paragraphs)
    option_paragraphs = [
        paragraph
        for paragraph in generated_doc.paragraphs
        if paragraph.text.strip().startswith(("A. ", "B. ", "C. ", "D. "))
    ]
    assert option_paragraphs
    assert all(paragraph.style.name != "List Bullet" for paragraph in option_paragraphs)
    assert all(not paragraph.text.startswith("A. A.") for paragraph in option_paragraphs)
    assert "HƯỚNG DẪN CHẤM VÀ ĐÁP ÁN" in paragraphs
    assert any(text.startswith("Mã đề: ") for text in paragraphs)
