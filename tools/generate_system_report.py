from __future__ import annotations

import ast
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = REPO_ROOT / "docs"
OUTPUT_PATH = OUTPUT_DIR / "Bao_cao_khoa_hoc_he_thong_AI_Personalized_Learning.docx"

ROUTER_PREFIXES: Dict[str, List[str]] = {
    "auth": ["/api/auth"],
    "classroom": ["/api/classroom"],
    "assessment": ["/api/assessment"],
    "upload": ["/api/upload"],
    "adaptive": ["/api/adaptive"],
    "stats": ["/api/stats"],
    "evaluation": ["/api/evaluation"],
    "admin": ["/api/admin"],
    "document": ["/api/documents"],
    "subject": ["/api/subjects"],
    "exam_generator": ["/api/exam"],
    "test_ocr": ["/api/test-ocr"],
    "teacher_agent": ["/api/teacher"],
    "orbit": ["/api/orbit"],
    "planning": ["/api/planning"],
    "notification": ["/api/notifications"],
    "debug": ["", "/api"],
}

MODULE_SUMMARIES: Dict[str, str] = {
    "auth": "xac thuc, phan quyen, quan ly phien dang nhap va thay doi mat khau",
    "classroom": "quan ly lop hoc, tham gia lop va danh sach thanh vien",
    "assessment": "sinh bai kiem tra, nop bai va luu lich su danh gia",
    "upload": "tai tai lieu, phan tich mon hoc va quan ly metadata tai lieu",
    "adaptive": "gia su AI ca nhan hoa va tom tat hoc lieu",
    "stats": "thong ke hoc tap ca nhan va theo lop",
    "evaluation": "chat danh gia hoc luc va phan tich muc do hoc tap",
    "admin": "quan tri nguoi dung he thong",
    "document": "thu vien tai lieu, preview, tai xuong va question bank theo tai lieu",
    "subject": "quan ly danh muc mon hoc",
    "exam_generator": "sinh de thi Word trong luong teacher/exam",
    "test_ocr": "sinh de OCR, tai file dap an va cham bai scan OMR",
    "teacher_agent": "Nova agent ho tro giang vien",
    "orbit": "Orbit agent dong hanh sinh vien va nhac nho dinh ky",
    "planning": "lap ke hoach hoc tap va dieu chinh lo trinh",
    "notification": "thong bao trong he thong",
    "debug": "kiem thu, theo doi va debug cac luong LLM",
}

TECH_STACK = [
    ("Backend", "FastAPI", "Xay dung REST API bat dong bo, sinh OpenAPI, to chuc router va middleware cho he thong."),
    ("Backend", "Uvicorn", "Dong vai tro ASGI server de phuc vu ung dung FastAPI trong moi truong phat trien va van hanh."),
    ("Backend", "Pydantic", "Rang buoc schema request/response, kiem tra du lieu dau vao va chuan hoa payload JSON."),
    ("Backend", "SQLAlchemy", "ORM chuyen doi giua doi tuong Python va bang SQL, ho tro query, quan he va transaction."),
    ("Database", "SQLite", "Co so du lieu quan he mac dinh cua du an; duoc cau hinh WAL, foreign key va busy timeout de tang do on dinh."),
    ("AI/RAG", "LangChain", "Cung cap lop tich hop cho loader tai lieu, vector store, embedding va luong truy hoi bo sung ngu canh."),
    ("AI/RAG", "ChromaDB", "Kho vector luu chunk tai lieu va embedding, phuc vu truy hoi ngu canh cho cac agent."),
    ("AI/RAG", "Sentence Transformers / all-MiniLM-L6-v2", "Sinh embedding cho chunk tai lieu de phuc vu tim kiem tuong dong ngu nghia."),
    ("AI/LLM", "Groq SDK", "Ket noi toi cac mo hinh LLM toc do cao, hien dang duoc dung boi adaptive agent, teacher agent, evaluation agent va planning agent."),
    ("AI/LLM", "Google Generative AI", "Kenh du phong cho mot so agent, dac biet assessment agent trong truong hop can doi provider."),
    ("Xu ly tai lieu", "PyPDFLoader, Docx2txtLoader, TextLoader, python-pptx", "Doc noi dung PDF, DOCX, TXT, PPTX de nap vao RAG va tao preview cho giao dien."),
    ("OCR", "PyMuPDF", "Raster hoa tung trang PDF scan thanh anh RGB de dua vao pipeline OMR."),
    ("OCR", "OpenCV", "Canh chinh marker, threshold, phat hien bubble, xu ly hinh hoc va crop cac vung thong tin."),
    ("OCR", "pytesseract", "Nhan dien chu viet/in trong vung ten sinh vien sau khi da crop bang image processing."),
    ("Sinh tai lieu", "python-docx", "Sinh file Word de thi va tai lieu bao cao he thong."),
    ("Sinh tai lieu", "Tu tao xlsx bang zip/XML", "He thong OCR tao workbook dap an theo tung ma de ma khong phu thuoc thu vien Excel nang."),
    ("Frontend", "Next.js", "Framework React cho giao dien web, to chuc route theo thu muc va ho tro SSR/CSR linh hoat."),
    ("Frontend", "React", "Xay dung UI theo component, quan ly state va luong tuong tac nguoi dung."),
    ("Frontend", "TypeScript", "Bo sung kieu tinh cho frontend de giam loi tich hop API."),
    ("Frontend", "Tailwind CSS", "Tao giao dien nhanh theo utility class, de dong bo phong cach giua cac man hinh."),
    ("Frontend", "Axios", "HTTP client cho frontend, ho tro timeout, upload multipart va xu ly loi thong nhat."),
    ("Frontend", "Recharts", "Ve bieu do thong ke tren dashboard va cac man hinh phan tich."),
    ("Frontend", "react-hot-toast", "Thong bao phan hoi nhanh cho nguoi dung khi tao de, upload, cham OCR, luu du lieu."),
    ("Van hanh", "RotatingFileHandler + request id", "Ghi log quay vong, gan ma yeu cau tren tung request va chuan hoa loi JSON cho truy vet."),
]

MODEL_PURPOSES = {
    "subjects": "Luu danh muc mon hoc va metadata mo ta cua tung mon.",
    "users": "Luu tai khoan, vai tro, thong tin nhan than va MSSV cua nguoi dung.",
    "classrooms": "Luu lop hoc, quan he voi mon hoc va giao vien phu trach.",
    "documents": "Luu tai lieu hoc tap duoc giao vien upload theo lop/mon.",
    "document_publications": "Luu trang thai co cho phep sinh vien nhin thay tai lieu hay khong.",
    "learning_roadmaps": "Luu lo trinh hoc tong quan do AI tao theo mon.",
    "learner_profiles": "Luu ho so nang luc tong hop theo mon cua sinh vien.",
    "study_sessions": "Luu phien hoc de thong ke effort va thoi luong hoc tap.",
    "assessment_history": "Luu ket qua cac bai danh gia va lich su lam bai.",
    "question_bank": "Luu ngan hang cau hoi theo mon, tai lieu, do kho va dap an.",
    "chunks": "Luu cac doan van ban da cat nho tu tai lieu de phuc vu RAG.",
    "assessment_results": "Luu ket qua danh gia rut gon theo chu de/mon hoc.",
    "student_learning_progress": "Luu tien do hoc tap tong hop va chi so hoat dong cua sinh vien.",
    "user_login_sessions": "Luu phien dang nhap/dang xuat de tinh muc do tham gia.",
    "orbit_chat_sessions": "Luu phien tro chuyen cua Orbit voi tung sinh vien.",
    "orbit_chat_messages": "Luu tung thong diep trong hoi thoai Orbit.",
    "orbit_coach_directives": "Luu chi dao cua giao vien giao cho Orbit de theo sat sinh vien.",
    "orbit_weekly_reminder_logs": "Luu lich su gui nhac nho hang tuan cua Orbit.",
    "student_document_evaluations": "Luu trang thai hoc va diem gan nhat tren tung tai lieu.",
    "student_document_score_history": "Luu lich su diem chi tiet theo tai lieu va moi lan danh gia.",
    "student_learning_plans": "Luu ban ke hoach hoc tap dang hoat dong cua sinh vien.",
    "student_learning_plan_steps": "Luu tung buoc/cong viec trong ke hoach hoc tap.",
    "notifications": "Luu thong bao noi bo cho giao vien va sinh vien.",
    "test_ocr_exam_batches": "Luu dot sinh de OCR, danh sach ma de, dap an va cau hinh OMR.",
    "test_ocr_grading_runs": "Luu moi lan upload PDF scan de cham OCR.",
    "test_ocr_grading_results": "Luu ket qua nhan dien tung phieu, dap an, diem so va debug OCR.",
}

AGENT_DESCRIPTIONS = [
    ("Nova Teacher Agent", "Tac nhan danh cho giao vien, giao tiep qua /api/teacher/*.", "Phan loai y dinh, nho ngu canh hoi thoai, tong hop thong tin lop/mon/sinh vien, dinh tuyen thao tac UI va co the kich hoat cac nghiep vu nhu tong quan lop, thong ke, tim tai lieu, sinh de thi hay tao chi dao Orbit."),
    ("Adaptive Agent", "Gia su AI ca nhan hoa theo hoc lieu va mon hoc.", "Su dung RAG tu ChromaDB va tai lieu da cong bo cho sinh vien de tra loi cau hoi, tom tat hoc lieu va de xuat roadmap hoc theo nang luc."),
    ("Planning Agent", "Tac nhan tao ke hoach hoc tap theo tai lieu va lich su hoc.", "Tu dong tao hoac tai tao ke hoach khi sinh vien dang nhap, sap xep thu tu hoc theo muc uu tien, cho phep chat dieu chinh ke hoach bang ngon ngu tu nhien."),
    ("Orbit Agent", "Tac nhan dong hanh sinh vien trong qua trinh hoc.", "Duy tri hoi thoai, de xuat tai lieu nen hoc tiep, theo doi tien do, tao lich su chat va phoi hop voi chi dao cua giao vien."),
    ("Evaluation Agent", "Tac nhan danh gia hoc luc/phan tich hoc tap.", "Doc du lieu lich su va thong ke de tra loi cau hoi phan tich danh gia, ho tro man hinh evaluation."),
    ("Assessment Agent", "Tac nhan sinh cau hoi va bai kiem tra.", "Tao question bank, tao quiz theo tai lieu, co fallback khi LLM/nguon cau hoi khong du."),
    ("Content Agent", "Tac nhan xu ly file dau vao cho RAG.", "Nhan dien mon hoc tu tai lieu, trich xuat noi dung, chia chunk va nap vector store."),
]

FRONTEND_PAGE_DESCRIPTIONS = {
    "/": ("cong khai", "Trang chu tong quan gioi thieu he thong va dieu huong vai tro."),
    "/auth": ("cong khai", "Dang nhap, dang ky va khoi tao thong tin phien lam viec tren localStorage."),
    "/adaptive": ("sinh vien", "Khong gian hoc ca nhan hoa, chat voi gia su AI va mo tai lieu dang hoc."),
    "/assessment": ("sinh vien", "Lam bai kiem tra va xem ket qua danh gia."),
    "/evaluation": ("giao vien/sinh vien", "Hoi dap phan tich hoc luc va thong ke danh gia."),
    "/planning": ("sinh vien", "Xem va dieu chinh ke hoach hoc tap do Planning Agent tao."),
    "/library": ("sinh vien", "Thu vien tai lieu da duoc cong bo cho lop hoc."),
    "/upload": ("giao vien", "Tai tai lieu hoc tap len he thong."),
    "/teacher": ("giao vien", "Dashboard giao vien, vao cac chuc nang Nova, tai lieu, lop, de thi."),
    "/teacher/subjects": ("giao vien", "Quan ly mon hoc va lop hoc."),
    "/teacher/documents": ("giao vien", "Quan ly thu vien tai lieu, question bank theo tai lieu."),
    "/teacher/question-bank": ("giao vien", "Xem, tao, sua, xoa ngan hang cau hoi."),
    "/teacher/members": ("giao vien", "Theo doi thanh vien lop va thong ke theo sinh vien."),
    "/teacher/exam": ("giao vien", "Luong sinh de Word thong thuong dang chay production."),
    "/test-ocr": ("giao vien", "Luong sinh de OCR, tai Word/Excel/PDF test va cham bai scan OMR."),
    "/admin/users": ("quan tri", "Quan ly tai khoan he thong."),
    "/admin/teachers": ("quan tri", "Tao giao vien moi."),
    "/admin/subjects": ("quan tri", "Dieu huong quan ly mon hoc danh cho admin."),
    "/admin/classrooms": ("quan tri", "Dieu huong quan ly lop hoc danh cho admin."),
    "/debug": ("ky su/quan tri", "Man hinh debug luong LLM va he thong."),
    "/debug1": ("ky su/quan tri", "Man hinh debug bo test rieng."),
    "/change-password": ("nguoi dung da dang nhap", "Thay doi mat khau."),
}

ENDPOINT_DESCRIPTIONS = {
    "register": "Dang ky tai khoan sinh vien moi.",
    "login": "Dang nhap, phat JWT va cap nhat login session.",
    "logout": "Dong phien dang nhap hien tai.",
    "change_password": "Doi mat khau cua nguoi dung.",
    "get_user_status": "Lay thong tin ho so va danh sach lop da tham gia.",
    "create_classroom": "Tao lop hoc moi.",
    "get_teacher_classes": "Lay danh sach lop do giao vien quan ly.",
    "get_all_classrooms": "Liet ke lop hoc theo bo loc.",
    "get_classroom": "Lay chi tiet lop hoc.",
    "update_classroom": "Cap nhat metadata lop hoc.",
    "delete_classroom": "Xoa lop hoc.",
    "join_classroom": "Sinh vien tham gia lop bang ma hoac thong tin lop.",
    "get_class_members": "Lay danh sach thanh vien lop.",
    "remove_student_from_class": "Xoa sinh vien khoi lop.",
    "analyze_document_subject": "Doan mon hoc phu hop cua tai lieu truoc khi upload.",
    "upload_document": "Nap tai lieu len server va xu ly RAG.",
    "get_documents": "Lay danh sach tai lieu theo giao vien/lop hoac theo thu vien lop.",
    "update_document_metadata": "Sua tieu de/metadata tai lieu.",
    "update_document_visibility": "Bat/tat cong bo tai lieu cho sinh vien.",
    "delete_document": "Xoa tai lieu va don dep du lieu lien quan.",
    "generate_question_bank_for_document": "Sinh question bank tu tai lieu.",
    "get_question_bank_for_document": "Lay cac cau hoi theo tai lieu.",
    "add_question_bank_for_document": "Them thu cong cau hoi vao question bank.",
    "update_question_bank_entry": "Sua cau hoi trong question bank.",
    "delete_question_bank_entry": "Xoa cau hoi khoi question bank.",
    "append_question_bank_for_document": "Bo sung them cau hoi vao question bank da co.",
    "get_student_documents": "Lay tai lieu sinh vien duoc quyen hoc.",
    "download_document": "Tai tai lieu goc.",
    "view_document_inline": "Mo tai lieu theo kieu inline tren trinh duyet.",
    "preview_document": "Trich xuat preview noi dung tai lieu.",
    "get_documents_by_subject": "Lay tai lieu theo mon hoc.",
    "generate_quiz": "Sinh de kiem tra thong thuong.",
    "generate_session_assessment": "Sinh bai danh gia theo phien hoc.",
    "ensure_document_question_bank": "Dam bao tai lieu da co question bank.",
    "generate_chapter_quiz": "Sinh quiz theo chuong/noi dung tai lieu.",
    "submit_quiz": "Nop bai va cham ket qua.",
    "save_quiz_result": "Luu ket qua quiz vao DB.",
    "get_learning_roadmap": "Lay lo trinh hoc theo mon.",
    "get_evaluation_history": "Lay lich su danh gia theo mon.",
    "get_learning_recommendation": "Lay goi y hoc tap tu adaptive agent.",
    "chat_with_adaptive_tutor": "Chat voi gia su AI.",
    "summarize_material": "Tom tat hoc lieu dang hoc.",
    "teacher_assistant": "Endpoint goi Nova cho giao vien.",
    "nova_interactive": "Endpoint Nova co action metadata cho frontend.",
    "chat_with_orbit": "Chat voi Orbit agent.",
    "get_orbit_history": "Lay lich su hoi thoai Orbit.",
    "get_orbit_progress": "Lay thong ke tien do hoc va tuong tac Orbit.",
    "create_teacher_directive": "Giao chi tieu hoc tap cua giao vien cho Orbit.",
    "get_weekly_inactivity_report": "Bao cao sinh vien it hoat dong theo tuan.",
    "trigger_weekly_reminders": "Gui nhac nho hang tuan tu Orbit.",
    "get_student_plan": "Lay ke hoach hoc tap hien tai.",
    "regenerate_plan": "Tai tao ke hoach hoc tap.",
    "planning_chat": "Dieu chinh ke hoach bang ngon ngu tu nhien.",
    "planning_chat_examples": "Lay mau cau lenh dieu chinh ke hoach.",
    "evaluation_chat": "Chat phan tich hoc luc.",
    "evaluation_chat_examples": "Lay vi du prompt cho evaluation.",
    "get_stats": "Thong ke hoc tap tong hop.",
    "get_class_analytics": "Thong ke theo lop hoc.",
    "create_user": "Tao tai khoan giao vien hoac sinh vien tu phia admin.",
    "delete_user": "Xoa tai khoan.",
    "get_all_users": "Liet ke nguoi dung he thong.",
    "list_subjects": "Liet ke mon hoc.",
    "get_subject": "Lay chi tiet mon hoc.",
    "create_subject": "Tao mon hoc.",
    "update_subject": "Sua mon hoc.",
    "delete_subject": "Xoa mon hoc.",
    "generate_exam_word": "Sinh file Word de thi trong luong teacher/exam.",
    "generate_test_ocr_word": "Sinh batch de OCR va file Word OCR.",
    "get_test_ocr_batch": "Lay metadata cua mot batch OCR.",
    "download_test_ocr_docx": "Tai file Word cua batch OCR.",
    "download_test_ocr_answer_xlsx": "Tai file Excel dap an OCR.",
    "download_test_ocr_test_sheet": "Tai phieu tra loi mau de test cham OCR.",
    "grade_test_ocr_pdf": "Upload PDF scan va cham OCR.",
    "update_test_ocr_student_name": "Cap nhat ten sinh vien da nhan dien/da chinh sua.",
    "list_notifications": "Lay thong bao cua nguoi dung.",
    "create_notification": "Tao thong bao moi.",
    "mark_notification_read": "Danh dau da doc mot thong bao.",
    "mark_all_notifications_read": "Danh dau da doc tat ca thong bao cua nguoi dung.",
    "debug_stream": "Mo kenh stream debug.",
    "debug_test": "Kiem tra endpoint debug co san sang khong.",
    "debug_health": "Kiem tra suc khoe he thong debug.",
    "debug_test_suite": "Chay bo test nghiep vu debug.",
    "debug_chat": "Thu nghiem hoi dap truc tiep voi LLM trong man hinh debug.",
}


@dataclass
class RouteInfo:
    module: str
    method: str
    route: str
    handler: str
    full_paths: List[str]


@dataclass
class ModelInfo:
    class_name: str
    table_name: str
    columns: List[str]


def normalize_path(prefix: str, route: str) -> str:
    if not prefix:
        return route or "/"
    if not route:
        return prefix
    return f"{prefix.rstrip('/')}/{route.lstrip('/')}"


def extract_routes() -> List[RouteInfo]:
    api_dir = REPO_ROOT / "backend" / "api"
    results: List[RouteInfo] = []
    for path in sorted(api_dir.glob("*.py")):
        module = path.stem
        source = path.read_text(encoding="utf-8")
        tree = ast.parse(source)
        prefixes = ROUTER_PREFIXES.get(module, [f"/api/{module}"])
        for node in tree.body:
            if not isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
                continue
            for decorator in node.decorator_list:
                if not isinstance(decorator, ast.Call):
                    continue
                if not isinstance(decorator.func, ast.Attribute):
                    continue
                if not isinstance(decorator.func.value, ast.Name) or decorator.func.value.id != "router":
                    continue
                route = ""
                if decorator.args and isinstance(decorator.args[0], ast.Constant):
                    route = str(decorator.args[0].value or "")
                method = decorator.func.attr.upper()
                full_paths = [normalize_path(prefix, route) for prefix in prefixes]
                results.append(RouteInfo(module=module, method=method, route=route, handler=node.name, full_paths=full_paths))
    return results


def _column_type_name(value: ast.AST) -> str:
    if isinstance(value, ast.Name):
        return value.id
    if isinstance(value, ast.Call):
        return _column_type_name(value.func)
    if isinstance(value, ast.Attribute):
        return value.attr
    return "Unknown"


def extract_models() -> List[ModelInfo]:
    source = (REPO_ROOT / "backend" / "db" / "models.py").read_text(encoding="utf-8")
    tree = ast.parse(source)
    results: List[ModelInfo] = []
    for node in tree.body:
        if not isinstance(node, ast.ClassDef):
            continue
        table_name = ""
        columns: List[str] = []
        for stmt in node.body:
            if not isinstance(stmt, ast.Assign):
                continue
            for target in stmt.targets:
                if not isinstance(target, ast.Name):
                    continue
                if target.id == "__tablename__":
                    if isinstance(stmt.value, ast.Constant):
                        table_name = str(stmt.value.value or "")
                    continue
                if isinstance(stmt.value, ast.Call) and _column_type_name(stmt.value.func) == "Column":
                    col_type = _column_type_name(stmt.value.args[0]) if stmt.value.args else "Unknown"
                    columns.append(f"{target.id} ({col_type})")
        if table_name:
            results.append(ModelInfo(class_name=node.name, table_name=table_name, columns=columns))
    return results


def extract_frontend_pages() -> List[tuple[str, str, str]]:
    app_dir = REPO_ROOT / "frontend" / "app"
    rows: List[tuple[str, str, str]] = []
    for page in sorted(app_dir.glob("**/page.tsx")):
        rel = page.relative_to(app_dir)
        if str(rel) == "page.tsx":
            route = "/"
        else:
            route = "/" + "/".join(rel.parts[:-1]).replace("\\", "/")
        audience, purpose = FRONTEND_PAGE_DESCRIPTIONS.get(route, ("noi bo", "Trang chuc nang cua he thong."))
        rows.append((route, audience, purpose))
    return rows


def set_default_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = "Bao cao mo ta he thong AI-Based Personalized Learning Platform - Trang "
        add_page_number(p)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def add_numbered(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)


def add_table(doc: Document, headers: Iterable[str], rows: Iterable[Iterable[str]]) -> None:
    headers = list(headers)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)


def write_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BAO CAO MO TA HE THONG\nAI-BASED PERSONALIZED LEARNING PLATFORM")
    r.bold = True
    r.font.size = Pt(18)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Phien ban tai lieu: tu dong tong hop tu ma nguon hien hanh\n").italic = True
    p2.add_run(f"Ngay tao: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    doc.add_paragraph(
        "Tai lieu nay mo ta toan dien he thong quan ly hoc tap ca nhan hoa tich hop AI Agent, "
        "RAG, sinh de thi va cham bai OCR. Noi dung duoc trich xuat truc tiep tu codebase backend, frontend, "
        "router API va schema co so du lieu dang ton tai trong du an."
    )
    doc.add_page_break()


def write_introduction(doc: Document, route_count: int, model_count: int, page_count: int) -> None:
    doc.add_heading("1. Tong quan va muc tieu tai lieu", level=1)
    doc.add_paragraph(
        "He thong AI-Based Personalized Learning Platform la mot nen tang ho tro hoc tap va quan ly hoc lieu "
        "theo huong ca nhan hoa. He thong khong chi cung cap giao dien web cho giao vien va sinh vien, ma con tich hop "
        "nhiều tac nhan AI de ho tro tu van hoc tap, lap ke hoach, thong ke, tao de va cham bai tu dong."
    )
    doc.add_paragraph(
        "Muc tieu cua tai lieu la mo ta khoa hoc kien truc, cong nghe, thanh phan backend/frontend, luong agent, "
        "luong OCR, danh muc API va thiet ke co so du lieu. Tai lieu dac biet tach rieng hai chuong trong tam: "
        "(i) he thong Agent va (ii) he thong tao de/cham thi OCR."
    )
    add_bullet(doc, f"He thong hien co it nhat {route_count} endpoint nghiep vu duoc dinh nghia trong thu muc backend/api.")
    add_bullet(doc, f"Schema du lieu co {model_count} model bang du lieu chinh (chua tinh bang trung gian enrollments).")
    add_bullet(doc, f"Frontend co {page_count} route giao dien theo co che file-based routing cua Next.js.")

    doc.add_heading("2. Pham vi va doi tuong phuc vu", level=1)
    add_bullet(doc, "Doi tuong quan tri: tao va quan ly tai khoan, mon hoc, lop hoc.")
    add_bullet(doc, "Doi tuong giao vien: quan ly lop, tai lieu, question bank, sinh de, su dung Nova va OCR.")
    add_bullet(doc, "Doi tuong sinh vien: hoc voi Adaptive/Orbit, xem tai lieu, thuc hien assessment, theo doi ke hoach hoc tap.")

    doc.add_heading("3. Phuong phap tong hop tai lieu", level=1)
    add_numbered(doc, "Phan tich ma nguon backend de trich xuat router, service, agent va cau hinh khoi dong.")
    add_numbered(doc, "Phan tich ma nguon frontend de tong hop route giao dien va cach tich hop API.")
    add_numbered(doc, "Phan tich models SQLAlchemy de lap bang schema du lieu.")
    add_numbered(doc, "He thong hoa cac thanh phan thanh mo hinh kien truc va cac luong nghiep vu chinh.")


def write_technology_section(doc: Document) -> None:
    doc.add_heading("4. Nen tang cong nghe", level=1)
    doc.add_paragraph(
        "He thong su dung kien truc web full-stack theo mo hinh tach lop ro rang. Moi cong nghe duoc lua chon de giai quyet "
        "một nhu cau cu the trong chuoi gia tri: phuc vu API, luu tru, RAG, AI Agent, xu ly tai lieu, OCR va giao dien nguoi dung."
    )
    add_table(doc, ["Nhom", "Cong nghe", "Vai tro khoa hoc - ky thuat"], TECH_STACK)

    doc.add_heading("5. Cau hinh va van hanh he thong", level=1)
    add_bullet(doc, "FastAPI app khoi dong voi lifespan, tu dong tao bang SQLAlchemy metadata, seed subject va admin mac dinh.")
    add_bullet(doc, "SQLite duoc cau hinh journal_mode=WAL, foreign_keys=ON, busy_timeout=30000 de tang do on dinh khi truy cap dong thoi.")
    add_bullet(doc, "RAG embeddings co the duoc bat/tat bang bien moi truong RAG_EMBEDDINGS_ENABLED; tren Windows mac dinh duoc tat de uu tien tinh on dinh.")
    add_bullet(doc, "Middleware logging tao request_id, ghi log app.log va error.log theo co che quay vong.")
    add_bullet(doc, "Du lieu tam va anh OCR duoc mount qua /temp_uploads de frontend co the preview ket qua crop.")


def write_architecture_section(doc: Document) -> None:
    doc.add_heading("6. Kien truc tong the", level=1)
    doc.add_paragraph(
        "He thong duoc thiet ke theo mo hinh nhieu lop, trong do presentation layer (Next.js) giao tiep doc quyen voi service layer "
        "(FastAPI). Business logic duoc tach thanh routers, services, agents, memory va rag modules. Du lieu duoc luu dong thoi tren "
        "co so du lieu quan he (SQLite/SQLAlchemy) va vector store (ChromaDB)."
    )
    add_bullet(doc, "Presentation layer: giao dien React/Next.js theo vai tro giao vien, sinh vien, admin.")
    add_bullet(doc, "API layer: cac router FastAPI trong backend/api, dinh nghia endpoint va validate request.")
    add_bullet(doc, "Service/Agent layer: nghiep vu AI, tao de, OCR, document preview, notification, reminder.")
    add_bullet(doc, "Persistence layer: SQLite cho du lieu cau truc; ChromaDB cho tri nho vector cua tai lieu.")

    doc.add_paragraph("So do logic tong the:")
    doc.add_paragraph(
        "Frontend (Next.js/React) -> FastAPI Router -> Service/Agent -> SQLAlchemy/SQLite + ChromaDB + File Storage (temp_uploads)",
    )

    doc.add_heading("7. Cau truc module backend", level=1)
    add_table(
        doc,
        ["Thu muc", "Chuc nang"],
        [
            ("backend/api", "Cac endpoint REST theo nghiep vu he thong."),
            ("backend/agents", "Tac nhan AI phuc vu giao vien, sinh vien, danh gia va sinh cau hoi."),
            ("backend/services", "Dich vu ho tro OCR, sinh de, PDF, thong ke, reminder, storage."),
            ("backend/db", "Cau hinh SQLAlchemy, session factory va models."),
            ("backend/rag", "Embedding va vector store cho truy hoi ngu canh."),
            ("backend/memory", "Bo nho hoi thoai, classifier y dinh va action router cua teacher agent."),
            ("temp_uploads", "Kho file tam, tai lieu da upload va ket qua OCR da raster/crop."),
        ],
    )

    doc.add_heading("8. Cau truc module frontend", level=1)
    doc.add_paragraph(
        "Frontend su dung App Router cua Next.js, trong do moi thu muc app/*/page.tsx tuong ung mot route. "
        "Du lieu phien nguoi dung va context chon lop/mon duoc luu chu yeu trong localStorage, trong khi axios/apiClient dam "
        "nhiem viec giao tiep HTTP voi backend."
    )


def write_frontend_section(doc: Document, pages: List[tuple[str, str, str]]) -> None:
    doc.add_heading("9. Danh muc giao dien frontend", level=1)
    add_table(doc, ["Route", "Doi tuong", "Muc dich"], pages)
    doc.add_paragraph(
        "Mot dac diem quan trong cua frontend la cac man hinh nghiep vu chuyen sau duoc tach rieng, "
        "vi du /teacher/exam cho luong sinh de thong thuong va /test-ocr cho luong OCR. "
        "Dieu nay giup giam rui ro anh huong tinh nang dang chay production."
    )


def write_feature_section(doc: Document) -> None:
    doc.add_heading("10. Chuc nang nghiep vu tong quat", level=1)
    add_numbered(doc, "Quan ly xac thuc va phan quyen: dang ky, dang nhap, JWT, doi mat khau, theo doi login session.")
    add_numbered(doc, "Quan ly mon hoc va lop hoc: tao mon, tao lop, tham gia lop, xem thanh vien.")
    add_numbered(doc, "Quan ly tai lieu va RAG: upload PDF/DOCX/PPTX/TXT, preview, cong bo cho sinh vien, nap chunk vao vector store.")
    add_numbered(doc, "Quan ly question bank va assessment: tao/sua/xoa ngan hang cau hoi, sinh quiz, nop bai, luu lich su ket qua.")
    add_numbered(doc, "Ho tro hoc tap bang Agent: Nova, Adaptive, Orbit, Planning, Evaluation.")
    add_numbered(doc, "Thong ke va canh bao: learning stats, class analytics, weekly inactivity report, notification.")
    add_numbered(doc, "Sinh de thi thong thuong: route /teacher/exam va backend /api/exam/generate-word.")
    add_numbered(doc, "Sinh de va cham thi OCR: route /test-ocr va backend /api/test-ocr/*.")


def write_agent_section(doc: Document) -> None:
    doc.add_heading("11. He thong Agent (chuong trong tam 1)", level=1)
    doc.add_paragraph(
        "He thong Agent la truc thong minh cua du an. Khac voi mo hinh chatbot don le, moi agent trong he thong phuc vu mot vai tro "
        "nghiep vu rieng, co du lieu bo tro, co bo nho hoi thoai hoac co rang buoc ve ngu canh hoc tap."
    )
    add_table(doc, ["Agent", "Pham vi", "Mo ta chuc nang"], AGENT_DESCRIPTIONS)

    doc.add_heading("11.1 Nova Teacher Agent", level=2)
    doc.add_paragraph(
        "Nova la tac nhan danh cho giao vien, duoc expose qua /api/teacher/assistant va /api/teacher/nova-interactive. "
        "Nova ket hop ba lop logic: bo phan nhan dien y dinh, bo nho hoi thoai va bo dinh tuyen hanh dong frontend."
    )
    add_bullet(doc, "Nhận diện ý định: teacher agent su dung IntentClassifier, co the ket hop rule-based va LLM de phan loai nhu class_overview, class_analytics, material, exam_generation.")
    add_bullet(doc, "Bo nho hoi thoai: memory.conversation_memory luu context theo cap teacher_id - class_id, giup agent xu ly follow-up va truy hoi ngữ cảnh gan nhat.")
    add_bullet(doc, "Action routing: memory.action_router sinh action_metadata de frontend biet nen mo man hinh nao, hien thanh phan nao hoac goi route nao tiep theo.")
    add_bullet(doc, "Nghiep vu giao vien: tong quan lop, tra cuu sinh vien, thong ke mon hoc, goi y tai lieu, sinh de thi, tao chi dao Orbit.")
    add_bullet(doc, "Kha nang du phong: endpoint nova-interactive co fallback path de tranh gãy luong khi xay ra loi noi bo.")

    doc.add_heading("11.2 Adaptive Agent", level=2)
    doc.add_paragraph(
        "Adaptive Agent la gia su AI phuc vu sinh vien. Agent nay doc du lieu hoc lieu cong bo, truy hoi chunk tu vector store "
        "va tong hop cau tra loi bang mo hinh Groq. Khi vector store hoac embedding khong kha dung, agent chuyen sang fallback mode "
        "de van cung cap huong dan hoc tap co ban."
    )
    add_bullet(doc, "Sinh roadmap hoc theo tai lieu va diem danh gia tren tung document.")
    add_bullet(doc, "Tom tat hoc lieu, giai thich khai niem va hoi dap theo tai lieu dang hoc.")
    add_bullet(doc, "Uu tien tai lieu duoc cong bo cho lop cua sinh vien, tranh truy xuat nham tai lieu ngoai pham vi.")

    doc.add_heading("11.3 Planning Agent", level=2)
    doc.add_paragraph(
        "Planning Agent la thanh phan bien tri thuc thanh ke hoach hanh dong. Agent nay duoc goi khi sinh vien dang nhap, khi "
        "frontend mo trang /planning va khi sinh vien yeu cau dieu chinh lich hoc."
    )
    add_bullet(doc, "Tao active plan cho sinh vien dua tren tai lieu, muc uu tien, deadline va lich su hoc.")
    add_bullet(doc, "Cho phep tai tao ke hoach thu cong hoac theo ly do login/manual.")
    add_bullet(doc, "Cho phep nhan cau lenh tu nhien nhu 'uu tien mon co so du lieu trong tuan nay' de sap xep lai step.")

    doc.add_heading("11.4 Orbit Agent", level=2)
    doc.add_paragraph(
        "Orbit la tro ly dong hanh cua sinh vien, khac voi Adaptive o cho no quan tam manh den nhịp hoc tap, lich su hoi thoai va "
        "su can thiep tu giao vien. Orbit co bang session/message rieng va co co che teacher directive."
    )
    add_bullet(doc, "Chat theo subject, class, document context va session history.")
    add_bullet(doc, "Goi y tai lieu nen mo tiep dua tren muc do tham gia va ket qua hoc.")
    add_bullet(doc, "Theo doi chi so tien do, so cau hoi Orbit, so phut chat, so bai hoc va bai test.")
    add_bullet(doc, "Sinh weekly inactivity report va gui nhac nho hang tuan qua services.orbit_reminders.")

    doc.add_heading("11.5 Evaluation va cac tac nhan phu tro", level=2)
    add_bullet(doc, "Evaluation Agent ho tro phan tich hoc luc, tra loi cau hoi danh gia va dien giai thong ke.")
    add_bullet(doc, "Assessment Agent sinh cau hoi, question bank va cac de kiem tra phuc vu danh gia hoc tap.")
    add_bullet(doc, "Content Agent tiep nhan tai lieu dau vao, xac dinh mon hoc, tach noi dung va bo sung tri thuc vao RAG.")


def write_ocr_section(doc: Document) -> None:
    doc.add_heading("12. He thong tao de va cham thi OCR (chuong trong tam 2)", level=1)
    doc.add_paragraph(
        "He thong OCR duoc phat trien tai route /test-ocr nhu mot luong doc lap voi /teacher/exam dang chay production. "
        "Muc tieu cua thiet ke nay la mo rong kha nang sinh de va cham bai hang loat ma khong xam pham logic dang van hanh cua teacher/exam."
    )

    doc.add_heading("12.1 Dau vao nghiep vu", level=2)
    add_bullet(doc, "Giao vien chon lop hoc, mon hoc, so cau hoi, so ma de, muc do va so cot MSSV.")
    add_bullet(doc, "He thong tai question bank va chunk tai lieu lien quan de tao de thi.")
    add_bullet(doc, "Moi batch OCR duoc gan mot batch_code va luu vao bang test_ocr_exam_batches.")

    doc.add_heading("12.2 Sinh de OCR", level=2)
    add_bullet(doc, "OCRExamGeneratorService tai su dung logic sinh de tu question bank va dap an.")
    add_bullet(doc, "He thong tao file Word de thi va dong thoi tao file Excel dap an theo tung sheet = tung ma de.")
    add_bullet(doc, "Khac voi ban sinh de thong thuong, batch OCR chi tao 1 phieu tra loi OMR chung; sinh vien tu to ma de tren phieu.")
    add_bullet(doc, "He thong tao them mot test sheet PDF duy nhat voi MSSV random va dap an random de kiem thu luong scan/cham.")

    doc.add_heading("12.3 Thiet ke phieu OMR", level=2)
    add_bullet(doc, "Kho giay A4, mau den-trang, toan bo toa do layout la co dinh.")
    add_bullet(doc, "Bon goc co alignment marker de phuc vu canh chinh va hieu chinh phep bien dang phoi canh.")
    add_bullet(doc, "Khu vuc Ho va ten la mot box rong; anh crop ten hien nay duoc tight crop de chi giu lai vung co chu.")
    add_bullet(doc, "Khu vuc MSSV la luoi bubble 0-9 theo tung cot, so cot co the cau hinh.")
    add_bullet(doc, "Khu vuc Ma de la luoi bubble rieng va khong duoc to san.")
    add_bullet(doc, "Khu vuc dap an phan bo dong theo so cau hoi, ho tro A/B/C/D va dieu chinh radius de OCR nhan bubble de hon.")

    doc.add_heading("12.4 File Excel dap an", level=2)
    add_bullet(doc, "Moi ma de duoc bieu dien boi mot worksheet rieng trong workbook.")
    add_bullet(doc, "Moi sheet chi luu hai cot question_number va correct_answer, trong do dap an duoc chuan hoa ve A/B/C/D.")
    add_bullet(doc, "Nguoi dung co the su dung ngay file nay de upload lai khi cham OCR, nham tach bo dap an khoi batch neu can chinh sua.")

    doc.add_heading("12.5 Pipeline cham OCR", level=2)
    add_numbered(doc, "Nhan file PDF scan tu frontend, moi trang duoc xem la mot phieu tra loi.")
    add_numbered(doc, "PDFProcessorService dung PyMuPDF de render tung trang thanh anh RGB 200 dpi.")
    add_numbered(doc, "OMRProcessorService tim alignment markers, hieu chinh perspective va resize ve template bounds.")
    add_numbered(doc, "He thong cat cac vung MSSV, ma de, dap an va crop vung ten sinh vien.")
    add_numbered(doc, "OpenCV threshold + contour + bubble fill ratio duoc dung de xac dinh so/Ma de/phuong an to kin.")
    add_numbered(doc, "pytesseract nhan dien ten sinh vien tren anh crop. Ket qua nay co the duoc nguoi dung sua va luu lai qua API patch.")
    add_numbered(doc, "He thong doi chieu exam code voi bo dap an batch hoac file Excel upload vao, tinh so cau dung va quy doi diem 10.")
    add_numbered(doc, "Ket qua tung phieu duoc luu vao test_ocr_grading_results va anh crop duoc expose qua /temp_uploads de frontend preview.")

    doc.add_heading("12.6 Trang thai va du lieu ket qua OCR", level=2)
    add_bullet(doc, "Trang thai co the gom: graded, missing_exam_code, missing_student_id, ambiguous_answers, unknown_exam_code.")
    add_bullet(doc, "Frontend hien thi anh ten SV, ten sinh vien co the sua, MSSV, ma de, dap an nhan dien, diem va trang thai.")
    add_bullet(doc, "Ten sinh vien do may OCR doc duoc va ten da nguoi dung sua deu duoc bao ton trong debug_json cua ket qua cham.")

    doc.add_heading("12.7 Gia tri ky thuat cua luong OCR", level=2)
    add_bullet(doc, "Toan bo nhan dien bubble mang tinh xac dinh (deterministic), khong dua vao mo hinh hoc may cho OMR.")
    add_bullet(doc, "Kien truc tach route /test-ocr giup giam rui ro doi voi nghiep vu /teacher/exam dang production.")
    add_bullet(doc, "Ket hop Word, Excel, PDF, OpenCV va Tesseract tao thanh mot chu trinh sinh de - in an - quet - cham diem khép kin.")


def write_load_balancing_section(doc: Document) -> None:
    doc.add_heading("13. Kien truc can bang tai va trien khai nhieu instance", level=1)
    doc.add_paragraph(
        "He thong da co san thiet ke production theo mo hinh scale ngang, trong do frontend va backend duoc nhan ban thanh "
        "nhieu instance va dat sau reverse proxy Nginx. Day la mot dac diem quan trong vi he thong chua cac luong tai nang nhu "
        "RAG, sinh de Word, OCR/OMR va chat Agent co do tre khong dong nhat."
    )

    doc.add_heading("13.1 Thanh phan can bang tai", level=2)
    add_table(
        doc,
        ["Lop", "Thanh phan", "Vai tro ky thuat"],
        [
            ("Reverse proxy", "Nginx", "Nhan request dau vao, dinh tuyen ve cum frontend/frontend_web hoac backend/backend_api, xu ly header X-Forwarded-* va websocket upgrade."),
            ("Frontend pool", "frontend_a, frontend_b", "Nhan ban ung dung Next.js thanh 2 instance de tang do san sang va phan tan tai truy cap giao dien."),
            ("Backend pool", "backend_a, backend_b", "Nhan ban API FastAPI thanh 2 instance de chia tai cho cac request AI, upload, OCR va thong ke."),
            ("App server", "Gunicorn + UvicornWorker", "Moi backend instance lai duoc xu ly boi nhieu worker/threads, tao thanh 2 tang song song hoa: giua instance va ben trong instance."),
            ("Shared state", "PostgreSQL, Redis, shared_temp_uploads, shared_chroma", "Dam bao cac instance backend dung chung nguon du lieu quan he, cache/queue, file tam va vector store."),
        ],
    )

    doc.add_heading("13.2 Chien luoc phan tai", level=2)
    add_bullet(doc, "Nginx dinh nghia upstream backend_api va frontend_web, moi upstream gom 2 server thanh vien.")
    add_bullet(doc, "Thuat toan least_conn duoc su dung de dua request moi toi instance dang co it ket noi dang xu ly nhat.")
    add_bullet(doc, "Route /api/* va /temp_uploads/* duoc proxy vao cum backend; route / duoc proxy vao cum frontend.")
    add_bullet(doc, "Cac header Upgrade/Connection duoc giu lai de ho tro cac luong ket noi keo dai va debug streaming.")
    add_bullet(doc, "proxy_read_timeout/proxy_send_timeout duoc nang len 300 giay cho cac request AI/OCR co the chay lau.")

    doc.add_heading("13.3 Kha nang san sang va readiness", level=2)
    add_bullet(doc, "Nginx uu tien route /api/ops/readiness de kiem tra tinh san sang cua backend truoc khi phuc vu tai nang.")
    add_bullet(doc, "backend/api/ops.py cung cap /health va /readiness; readiness tra 503 neu mot thanh phan co trang thai khong ok.")
    add_bullet(doc, "services/system_health.py kiem tra database latency, trang thai connection pool, conversation memory va tuy chon vector store.")
    add_bullet(doc, "Trong Docker Compose production, postgres va redis deu co healthcheck rieng truoc khi backend khoi dong.")

    doc.add_heading("13.4 Song song hoa trong moi backend instance", level=2)
    add_bullet(doc, "Gunicorn bind mac dinh tai 0.0.0.0:8010 va su dung UvicornWorker de chay FastAPI theo chuan ASGI.")
    add_bullet(doc, "So worker mac dinh duoc tinh theo CPU, co the override bang GUNICORN_WORKERS; threads duoc cau hinh bang GUNICORN_THREADS.")
    add_bullet(doc, "Cac tham so max_requests va max_requests_jitter giup quay vong worker chu dong, giam nguy co suy hao bo nho tren tai dai han.")
    add_bullet(doc, "Tang SQLAlchemy pool duoc cau hinh voi pool_size, max_overflow, pool_timeout, pool_recycle va pool_pre_ping de phuc vu truy cap dong thoi on dinh.")

    doc.add_heading("13.5 Y nghia kien truc", level=2)
    add_bullet(doc, "Kien truc nay cho phep he thong dung duoc tai hon hop: request nhe (CRUD, auth) va request nang (OCR, sinh de, agent chat).")
    add_bullet(doc, "Viec tach backend thanh nhieu instance giam nguy co mot request cham lam nghen toan bo dich vu.")
    add_bullet(doc, "Vi Redis, PostgreSQL, file tam va ChromaDB duoc chia se, nguoi dung co the duoc phuc vu boi bat ky backend instance nao ma van giu tinh nhat quan nghiep vu.")
    add_bullet(doc, "Mo hinh nay la nen tang phu hop de mo rong them autoscaling, queue hoa tac vu OCR hoac tach rieng cum AI ve sau.")



def write_api_section(doc: Document, routes: List[RouteInfo]) -> None:
    doc.add_heading("14. Danh muc API", level=1)
    doc.add_paragraph(
        "Bang duoi day liet ke cac endpoint nghiep vu duoc trich xuat truc tiep tu thu muc backend/api. "
        "Voi module debug, route duoc mount dong thoi tai root va /api."
    )
    modules = sorted({route.module for route in routes})
    for module in modules:
        module_routes = [route for route in routes if route.module == module]
        doc.add_heading(f"14.{modules.index(module) + 1}. Module {module}", level=2)
        doc.add_paragraph(f"Mo ta module: {MODULE_SUMMARIES.get(module, 'Nghiep vu chua duoc mo ta bo sung')}.")
        rows = []
        for route in module_routes:
            rows.append(
                (
                    route.method,
                    "\n".join(route.full_paths),
                    route.handler,
                    ENDPOINT_DESCRIPTIONS.get(route.handler, f"Endpoint thuoc module {module}."),
                )
            )
        add_table(doc, ["Method", "Duong dan", "Handler", "Muc dich"], rows)


def write_database_section(doc: Document, models: List[ModelInfo]) -> None:
    doc.add_heading("15. Thiet ke co so du lieu", level=1)
    doc.add_paragraph(
        "Tang du lieu cua he thong duoc trien khai bang SQLAlchemy ORM tren nen SQLite mac dinh. "
        "Schema bao gom cac bang phuc vu danh muc hoc tap, tai lieu, assessment, ke hoach hoc, tuong tac agent, thong bao va OCR."
    )
    add_bullet(doc, "Bang trung gian enrollments bieu dien quan he nhieu-nhieu giua users va classrooms.")
    add_bullet(doc, "Cac truong JSON duoc dung de luu cau truc co do dong cao nhu roadmap_data, answer_key_json, omr_layout_json, debug_json.")

    for index, model in enumerate(models, start=1):
        doc.add_heading(f"15.{index}. Bang {model.table_name}", level=2)
        doc.add_paragraph(
            MODEL_PURPOSES.get(model.table_name, f"Bang {model.table_name} luu du lieu nghiep vu cua model {model.class_name}.")
        )
        chunks: List[str] = []
        current = ""
        for column in model.columns:
            item = column
            if len(current) + len(item) + 2 > 1000:
                chunks.append(current.rstrip(", "))
                current = ""
            current += item + ", "
        if current:
            chunks.append(current.rstrip(", "))
        for chunk in chunks:
            doc.add_paragraph(f"Cot chinh: {chunk}")

    doc.add_heading("16. Nhom bang theo mien nghiep vu", level=1)
    add_table(
        doc,
        ["Mien du lieu", "Bang dai dien"],
        [
            ("Danh muc va nguoi dung", "subjects, users, classrooms, enrollments"),
            ("Tai lieu va RAG", "documents, document_publications, chunks, question_bank"),
            ("Danh gia hoc tap", "assessment_history, assessment_results, student_document_evaluations, student_document_score_history"),
            ("Theo doi tien do", "study_sessions, student_learning_progress, user_login_sessions"),
            ("Lo trinh va ke hoach", "learning_roadmaps, learner_profiles, student_learning_plans, student_learning_plan_steps"),
            ("Agent va thong bao", "orbit_chat_sessions, orbit_chat_messages, orbit_coach_directives, orbit_weekly_reminder_logs, notifications"),
            ("OCR", "test_ocr_exam_batches, test_ocr_grading_runs, test_ocr_grading_results"),
        ],
    )


def write_quality_section(doc: Document) -> None:
    doc.add_heading("17. Chat luong phan mem, logging va kha nang mo rong", level=1)
    add_bullet(doc, "RequestLoggingMiddleware gan request_id cho moi request, giup truy vet loi xuyen suot BE-FE.")
    add_bullet(doc, "error_json_response chuan hoa payload loi, bao gom detail, request_id va retryable.")
    add_bullet(doc, "Question bank warmup co the chay nen khi bat bien moi truong ENABLE_QUESTION_BANK_WARMUP.")
    add_bullet(doc, "Cau truc tach module giup he thong de mo rong them agent moi hoac route nghiep vu moi ma khong can pha vo toan bo codebase.")
    add_bullet(doc, "Luong OCR duoc tach khoi /teacher/exam, the hien mot quy tac thiet ke an toan cho production.")

    doc.add_heading("18. Ket luan", level=1)
    doc.add_paragraph(
        "Xet tren goc nhin khoa hoc he thong, day la mot nen tang tich hop da thanh phan: "
        "co so du lieu quan he, vector database, agent dung LLM, quan ly tai lieu, giao dien web va OCR xac dinh. "
        "Gia tri cua du an nam o kha nang ket noi du lieu hoc tap, hanh vi hoc tap va cong cu AI thanh mot he thong "
        "dong bo phuc vu ca giao vien lan sinh vien."
    )
    doc.add_paragraph(
        "Hai cum tinh nang trong tam la he thong Agent va he thong OCR da duoc thiet ke theo huong tach lop, "
        "giup giam ket dinh, tang kha nang bao tri va cho phep mo rong nghiep vu trong tuong lai. "
        "Tai lieu nay co the duoc su dung nhu mot ban mo ta he thong phuc vu bao cao khoa hoc, nghiem thu ky thuat, "
        "hoac chuyen giao cho nhom phat trien tiep theo."
    )


def build_report() -> Path:
    routes = extract_routes()
    models = extract_models()
    pages = extract_frontend_pages()

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_default_style(doc)
    add_footer(doc)
    write_cover(doc)
    write_introduction(doc, len(routes), len(models), len(pages))
    write_technology_section(doc)
    write_architecture_section(doc)
    write_frontend_section(doc, pages)
    write_feature_section(doc)
    write_agent_section(doc)
    write_ocr_section(doc)
    write_load_balancing_section(doc)
    doc.add_section(WD_SECTION.NEW_PAGE)
    write_api_section(doc, routes)
    doc.add_section(WD_SECTION.NEW_PAGE)
    write_database_section(doc, models)
    write_quality_section(doc)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    output = build_report()
    print(output)
