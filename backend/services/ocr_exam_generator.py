from __future__ import annotations

import copy
import io
import json
import math
import random
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Mm, Pt
from PIL import Image, ImageDraw, ImageFont
from sqlalchemy.orm import Session

from db import models
from services.exam_doc_utils import (
    DEFAULT_DEPARTMENT_NAME,
    DEFAULT_EXAM_TITLE,
    DEFAULT_EXAM_TYPE_LABEL,
    DEFAULT_SCHOOL_NAME,
    add_answer_key_section,
    add_exam_info_table,
    add_exam_instruction_block,
    add_multiple_choice_questions,
    add_school_header,
    apply_document_style,
)
from services.test_ocr_answer_key_excel import build_answer_key_workbook
from services.test_ocr_storage import (
    build_batch_code,
    build_generated_answer_xlsx_path,
    build_generated_docx_path,
    default_omr_config,
    ensure_test_ocr_dirs,
)


def _strip_option_prefix(value: str) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    text = re.sub(r"^[\u2022\-\*]+\s*", "", text)
    text = re.sub(r"^(?:[A-Da-d])[\.\)\:\-]\s*", "", text)
    text = re.sub(r"\s+#\d+\b", "", text)
    return text.strip()


def _normalize_question_bank_options(raw_options) -> List[str]:
    if isinstance(raw_options, list):
        return [_strip_option_prefix(str(item or "")) for item in raw_options if _strip_option_prefix(str(item or ""))]
    if isinstance(raw_options, str):
        try:
            parsed = json.loads(raw_options)
        except Exception:
            parsed = None
        if isinstance(parsed, list):
            return [_strip_option_prefix(str(item or "")) for item in parsed if _strip_option_prefix(str(item or ""))]
    return []


def _resolve_correct_label(correct_answer: str, options: List[str]) -> str:
    normalized = str(correct_answer or "").strip().upper()
    if normalized in {"A", "B", "C", "D"}:
        return normalized

    clean_answer = re.sub(r"\s+", " ", str(correct_answer or "")).strip().lower()
    for idx, option in enumerate(options[:4]):
        if re.sub(r"\s+", " ", str(option or "")).strip().lower() == clean_answer:
            return ["A", "B", "C", "D"][idx]
    return "A"


def _collect_chunk_texts(db: Session, allowed_filenames: List[str], limit: int = 200) -> List[str]:
    if not allowed_filenames:
        return []
    rows = (
        db.query(models.Chunk)
        .filter(models.Chunk.source_file.in_(allowed_filenames))
        .order_by(models.Chunk.id.asc())
        .limit(limit)
        .all()
    )
    return [re.sub(r"\s+", " ", str(row.content or "")).strip() for row in rows if str(row.content or "").strip()]


def _build_fallback_exam_questions(subject: str, exam_type: str, num_questions: int, context_summary: str, chunk_texts: List[str]) -> List[dict]:
    needed = max(1, int(num_questions or 1))
    source_text = "\n".join(chunk_texts[:120]) if chunk_texts else context_summary
    source_text = re.sub(r"\s+", " ", source_text or "").strip()

    sentences = []
    for item in re.split(r"(?<=[\.\!\?\;\:])\s+", source_text):
        clean = re.sub(r"\s+", " ", item or "").strip()
        if 45 <= len(clean) <= 220:
            sentences.append(clean)
    if not sentences:
        sentences = [f"Môn {subject} yêu cầu nắm chắc các khái niệm và nguyên lý cốt lõi."]

    result: List[dict] = []
    if exam_type == "trac_nghiem":
        for idx in range(needed):
            sentence = sentences[idx % len(sentences)]
            result.append(
                {
                    "q": f"Theo học liệu môn {subject}, phát biểu nào đúng nhất với ý sau: \"{sentence[:120]}\"?",
                    "options": [
                        sentence[:180],
                        (sentence[:120] + " nhưng áp dụng trong ngữ cảnh khác của môn học.")[:180],
                        f"Nội dung này không liên quan trực tiếp đến trọng tâm của {subject}.",
                        f"Đây chỉ là một ví dụ phụ, không phải kết luận chính về {subject}.",
                    ],
                    "ans": "A",
                    "exp": "Câu bổ sung tự động khi ngân hàng câu hỏi chưa đủ.",
                }
            )
    return result


def _remove_accents(value: str) -> str:
    safe = str(value or "").replace("đ", "d").replace("Đ", "D")
    nfkd_form = unicodedata.normalize("NFKD", safe)
    return "".join([char for char in nfkd_form if not unicodedata.combining(char)])


def _safe_filename_fragment(value: str) -> str:
    ascii_value = _remove_accents(value or "")
    ascii_value = re.sub(r"[^A-Za-z0-9_-]+", "_", ascii_value).strip("_")
    return ascii_value or "TaiLieu"


def _normalize_exam_type_key(value: str) -> str:
    raw = str(value or "").strip().lower()
    normalized = _remove_accents(raw).lower().replace("-", " ").replace("_", " ")
    compact = re.sub(r"\s+", " ", normalized).strip()
    ascii_only = re.sub(r"[^a-z]+", " ", compact).strip()

    candidates = {compact.replace(" ", "_"), ascii_only.replace(" ", "_")}
    if {"trac_nghiem", "tracnghiem"} & candidates:
        return "trac_nghiem"
    if compact.startswith("tr") and "nghi" in compact:
        return "trac_nghiem"
    if ascii_only.startswith("tr") and "nghi" in ascii_only:
        return "trac_nghiem"
    return candidates.pop() if candidates else ""


def _load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/timesbd.ttf" if bold else "C:/Windows/Fonts/times.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        "arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except Exception:
            continue
    return ImageFont.load_default()


def _draw_bubble(draw: ImageDraw.ImageDraw, center_x: int, center_y: int, radius: int, filled: bool = False) -> None:
    bbox = (center_x - radius, center_y - radius, center_x + radius, center_y + radius)
    draw.ellipse(bbox, outline="black", width=3, fill="black" if filled else "white")


def _draw_alignment_markers(draw: ImageDraw.ImageDraw, width: int, height: int, size: int, margin: int) -> None:
    for x, y in (
        (margin, margin),
        (width - margin - size, margin),
        (margin, height - margin - size),
        (width - margin - size, height - margin - size),
    ):
        draw.rectangle((x, y, x + size, y + size), fill="black", outline="black")


def _build_omr_layout(question_count: int, student_id_columns: int, exam_code_columns: int) -> Dict[str, object]:
    template = default_omr_config(
        question_count=question_count,
        student_id_columns=student_id_columns,
        exam_code_columns=exam_code_columns,
    )
    width = int(template["template_width"])
    height = int(template["template_height"])

    if question_count <= 20:
        answer_columns = 2
    elif question_count <= 45:
        answer_columns = 3
    else:
        answer_columns = 4
    questions_per_column = math.ceil(question_count / answer_columns)

    if questions_per_column <= 15:
        answer_bubble_radius = 30
    elif questions_per_column <= 20:
        answer_bubble_radius = 28
    elif questions_per_column <= 25:
        answer_bubble_radius = 24
    elif questions_per_column <= 30:
        answer_bubble_radius = 22
    else:
        answer_bubble_radius = 20

    outer_frame = {"x": 190, "y": 190, "w": width - 380, "h": height - 380}
    title_box = {"x": 180, "y": 180, "w": width - 360, "h": 210}
    name_box = {"x": 200, "y": 430, "w": 2080, "h": 150}
    student_id_block = {"x": 200, "y": 640, "w": 1180, "h": 690}
    exam_code_block = {"x": 1450, "y": 640, "w": 830, "h": 690}

    digit_row_top = student_id_block["y"] + 120
    digit_row_gap = 58
    digit_bubble_radius = 28
    digit_row_centers = [digit_row_top + idx * digit_row_gap for idx in range(10)]

    student_id_group_left = student_id_block["x"] + 150
    student_id_gap = 96
    student_id_x_centers = [student_id_group_left + idx * student_id_gap for idx in range(student_id_columns)]

    exam_code_group_left = exam_code_block["x"] + 150
    exam_code_gap = 108
    exam_code_x_centers = [exam_code_group_left + idx * exam_code_gap for idx in range(exam_code_columns)]

    answer_panel = {"x": 200, "y": 1390, "w": 2080, "h": 1840}
    answer_area_left = answer_panel["x"] + 18
    answer_area_right = answer_panel["x"] + answer_panel["w"] - 18
    answer_area_top = answer_panel["y"] + 210
    answer_area_bottom = answer_panel["y"] + answer_panel["h"] - 95
    column_gutter = 32
    total_gutter = column_gutter * (answer_columns - 1)
    block_width = (answer_area_right - answer_area_left - total_gutter) // answer_columns
    available_height = answer_area_bottom - answer_area_top
    answer_row_gap = max(
        answer_bubble_radius * 2 + 10,
        min(84, int(available_height / max(questions_per_column, 1))),
    )

    answer_columns_meta: List[Dict[str, object]] = []
    for column_index in range(answer_columns):
        base_x = answer_area_left + column_index * (block_width + column_gutter)
        option_group_width = min(300, max(210, block_width - 230))
        option_left = base_x + 138
        option_step = option_group_width / 3.0
        option_centers = [int(round(option_left + option_step * idx)) for idx in range(4)]

        question_start = column_index * questions_per_column + 1
        question_end = min(question_count, question_start + questions_per_column - 1)
        question_rows = [
            {"question_number": question_number, "y": answer_area_top + local_row_index * answer_row_gap}
            for local_row_index, question_number in enumerate(range(question_start, question_end + 1))
        ]

        answer_columns_meta.append(
            {
                "base_x": base_x,
                "block_width": block_width,
                "label_x": base_x + 10,
                "option_centers": option_centers,
                "header_y": answer_panel["y"] + 128,
                "question_rows": question_rows,
            }
        )

    return {
        **template,
        "template_bounds": {"width": width, "height": height},
        "outer_frame": outer_frame,
        "title_box": title_box,
        "name_box": name_box,
        "student_id_block": student_id_block,
        "exam_code_block": exam_code_block,
        "answer_panel": answer_panel,
        "alignment_markers": {"size": 72, "margin": 96},
        "digit_row_centers": digit_row_centers,
        "student_id_x_centers": student_id_x_centers,
        "exam_code_x_centers": exam_code_x_centers,
        "student_id_digit_label_x": student_id_block["x"] + 54,
        "exam_code_digit_label_x": exam_code_block["x"] + 54,
        "digit_bubble_radius": digit_bubble_radius,
        "answer_bubble_radius": answer_bubble_radius,
        "answer_columns_meta": answer_columns_meta,
        "answer_columns": answer_columns,
        "questions_per_column": questions_per_column,
        "instructions": [
            "Tô đậm kín một ô cho mỗi câu hỏi.",
            "Không để nét bút chạm vào viền ô tròn.",
        ],
    }


def _render_omr_sheet(
    version_code: str | None,
    question_count: int,
    layout: Dict[str, object],
    *,
    prefilled_student_id: str | None = None,
    prefilled_answers: Optional[List[str]] = None,
    school_name: str = DEFAULT_SCHOOL_NAME,
    department_name: str = DEFAULT_DEPARTMENT_NAME,
    subject_name: str = "",
    class_name: str = "",
    exam_date: str = "....... / ....... / ............",
) -> io.BytesIO:
    template_bounds = layout["template_bounds"]
    width = int(template_bounds["width"])
    height = int(template_bounds["height"])
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    title_font = _load_font(42, bold=True)
    subtitle_font = _load_font(28, bold=True)
    text_font = _load_font(24)
    small_font = _load_font(20)
    answer_label_font = _load_font(24, bold=True)
    question_font = _load_font(24, bold=True)
    header_font = _load_font(20, bold=True)
    meta_font = _load_font(18)

    marker_cfg = layout["alignment_markers"]
    _draw_alignment_markers(draw, width, height, int(marker_cfg["size"]), int(marker_cfg["margin"]))

    outer_frame = layout["outer_frame"]
    draw.rectangle(
        (outer_frame["x"], outer_frame["y"], outer_frame["x"] + outer_frame["w"], outer_frame["y"] + outer_frame["h"]),
        outline="black",
        width=2,
    )

    title_box = layout["title_box"]
    draw.text((width // 2, title_box["y"] + 16), school_name, fill="black", font=header_font, anchor="ma")
    draw.text((width // 2, title_box["y"] + 42), department_name, fill="black", font=header_font, anchor="ma")
    draw.text((width // 2, title_box["y"] + 90), "PHIẾU TRẢ LỜI TRẮC NGHIỆM", fill="black", font=title_font, anchor="ma")
    draw.text(
        (width // 2, title_box["y"] + 130),
        f"Môn học: {subject_name or '................................'}    Lớp: {class_name or '................................'}",
        fill="black",
        font=meta_font,
        anchor="ma",
    )
    draw.text(
        (width // 2, title_box["y"] + 156),
        f"Ngày thi: {exam_date}    Mã đề: tô tại khung bên phải",
        fill="black",
        font=meta_font,
        anchor="ma",
    )

    name_box = layout["name_box"]
    draw.text((name_box["x"], name_box["y"] - 42), "Họ và tên", fill="black", font=subtitle_font)
    draw.rectangle((name_box["x"], name_box["y"], name_box["x"] + name_box["w"], name_box["y"] + name_box["h"]), outline="black", width=4)
    draw.rectangle(
        (name_box["x"] + 18, name_box["y"] + 18, name_box["x"] + name_box["w"] - 18, name_box["y"] + name_box["h"] - 18),
        outline="black",
        width=1,
    )

    digit_bubble_radius = int(layout["digit_bubble_radius"])
    digit_row_centers = [int(value) for value in layout["digit_row_centers"]]

    student_id_block = layout["student_id_block"]
    draw.rectangle(
        (student_id_block["x"], student_id_block["y"], student_id_block["x"] + student_id_block["w"], student_id_block["y"] + student_id_block["h"]),
        outline="black",
        width=2,
    )
    student_id_x_centers = [int(value) for value in layout["student_id_x_centers"]]
    draw.text((student_id_block["x"] + 22, student_id_block["y"] + 24), "Mã sinh viên", fill="black", font=subtitle_font)
    student_id_label_x = int(layout["student_id_digit_label_x"])
    student_id_digits = list(str(prefilled_student_id or "").strip())
    for column_index, center_x in enumerate(student_id_x_centers, start=1):
        draw.text((center_x, student_id_block["y"] + 90), str(column_index), fill="black", font=small_font, anchor="ms")
    for digit, center_y in enumerate(digit_row_centers):
        draw.text((student_id_label_x, center_y), str(digit), fill="black", font=text_font, anchor="mm")
        for column_index, center_x in enumerate(student_id_x_centers):
            is_filled = column_index < len(student_id_digits) and student_id_digits[column_index].isdigit() and int(student_id_digits[column_index]) == digit
            _draw_bubble(draw, center_x, center_y, digit_bubble_radius, filled=is_filled)

    exam_code_block = layout["exam_code_block"]
    draw.rectangle(
        (exam_code_block["x"], exam_code_block["y"], exam_code_block["x"] + exam_code_block["w"], exam_code_block["y"] + exam_code_block["h"]),
        outline="black",
        width=2,
    )
    exam_code_x_centers = [int(value) for value in layout["exam_code_x_centers"]]
    draw.text((exam_code_block["x"] + 22, exam_code_block["y"] + 24), "Mã đề", fill="black", font=subtitle_font)
    exam_code_label_x = int(layout["exam_code_digit_label_x"])
    code_digits = list(str(version_code or "").strip())
    for column_index, center_x in enumerate(exam_code_x_centers, start=1):
        draw.text((center_x, exam_code_block["y"] + 90), str(column_index), fill="black", font=small_font, anchor="ms")
    for digit, center_y in enumerate(digit_row_centers):
        draw.text((exam_code_label_x, center_y), str(digit), fill="black", font=text_font, anchor="mm")
        for column_index, center_x in enumerate(exam_code_x_centers, start=1):
            filled_digit = int(code_digits[column_index - 1]) if column_index - 1 < len(code_digits) and code_digits[column_index - 1].isdigit() else None
            _draw_bubble(draw, center_x, center_y, digit_bubble_radius, filled=filled_digit == digit)

    answer_panel = layout["answer_panel"]
    answer_bubble_radius = int(layout["answer_bubble_radius"])
    answer_columns_meta = list(layout["answer_columns_meta"])
    draw.rectangle(
        (answer_panel["x"], answer_panel["y"], answer_panel["x"] + answer_panel["w"], answer_panel["y"] + answer_panel["h"]),
        outline="black",
        width=2,
    )
    draw.text((answer_panel["x"] + 22, answer_panel["y"] + 34), "Đánh dấu đáp án", fill="black", font=subtitle_font)
    draw.line(
        (answer_panel["x"] + 18, answer_panel["y"] + 92, answer_panel["x"] + answer_panel["w"] - 18, answer_panel["y"] + 92),
        fill="black",
        width=2,
    )

    option_labels = ["A", "B", "C", "D"]
    answer_index_by_label = {"A": 0, "B": 1, "C": 2, "D": 3}
    normalized_prefilled_answers = [str(item or "").strip().upper() for item in (prefilled_answers or [])]

    for column_meta in answer_columns_meta:
        base_x = int(column_meta["base_x"])
        block_width = int(column_meta["block_width"])
        header_y = int(column_meta["header_y"])
        draw.rectangle(
            (base_x - 6, answer_panel["y"] + 112, base_x + block_width + 8, answer_panel["y"] + answer_panel["h"] - 24),
            outline="black",
            width=1,
        )
        for option_index, center_x in enumerate(column_meta["option_centers"]):
            draw.text((int(center_x), header_y), option_labels[option_index], fill="black", font=answer_label_font, anchor="ms")
        for row_meta in column_meta["question_rows"]:
            center_y = int(row_meta["y"])
            question_number = int(row_meta["question_number"])
            draw.text((int(column_meta["label_x"]), center_y - 14), f"{question_number:02d}", fill="black", font=question_font)
            marked_answer = normalized_prefilled_answers[question_number - 1] if question_number - 1 < len(normalized_prefilled_answers) else ""
            marked_index = answer_index_by_label.get(marked_answer, -1)
            for option_index, center_x in enumerate(column_meta["option_centers"]):
                _draw_bubble(draw, int(center_x), center_y, answer_bubble_radius, filled=option_index == marked_index)

    stream = io.BytesIO()
    image.save(stream, format="PNG")
    stream.seek(0)
    return stream


@dataclass
class GeneratedBatchBundle:
    batch: models.TestOCRExamBatch
    file_path: Path
    download_filename: str
    answer_key_file_path: Path
    answer_key_download_filename: str


class OCRExamGeneratorService:
    def __init__(self, db: Session):
        self.db = db

    def _build_question_pool(
        self,
        *,
        subject: str,
        exam_type: str,
        num_questions: int,
        target_class: models.Classroom,
        docs: List[models.Document],
        context_summary: str,
        chunk_texts: List[str],
    ) -> List[dict]:
        allowed_filenames = {str(doc.filename or "").strip() for doc in docs if str(doc.filename or "").strip()}
        subject_id = int(target_class.subject_id)

        rows = (
            self.db.query(models.QuestionBank)
            .filter(
                models.QuestionBank.subject_id == subject_id,
                models.QuestionBank.source_file.in_(allowed_filenames),
            )
            .order_by(models.QuestionBank.id.asc())
            .all()
        )
        if len(rows) < num_questions:
            fallback_rows = (
                self.db.query(models.QuestionBank)
                .filter(models.QuestionBank.subject_id == subject_id)
                .order_by(models.QuestionBank.id.asc())
                .all()
            )
            seen_ids = {int(item.id) for item in rows}
            rows.extend([row for row in fallback_rows if int(row.id) not in seen_ids])

        question_pool: List[dict] = []
        for row in rows:
            options = _normalize_question_bank_options(getattr(row, "options", []))
            if len(options) < 4:
                continue
            question_text = re.sub(r"\s+", " ", str(getattr(row, "content", "") or "")).strip()
            if not question_text:
                continue
            question_pool.append(
                {
                    "q": question_text,
                    "options": options[:4],
                    "ans": _resolve_correct_label(str(getattr(row, "correct_answer", "") or ""), options),
                    "exp": re.sub(r"\s+", " ", str(getattr(row, "explanation", "") or "")).strip(),
                }
            )

        if len(question_pool) < num_questions:
            question_pool.extend(
                _build_fallback_exam_questions(
                    subject=subject,
                    exam_type=exam_type,
                    num_questions=num_questions - len(question_pool),
                    context_summary=context_summary,
                    chunk_texts=chunk_texts,
                )
            )
        return question_pool

    def _prepare_section(self, doc: Document) -> None:
        section = doc.sections[-1]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(12.7)
        section.bottom_margin = Mm(12.7)
        section.left_margin = Mm(12.7)
        section.right_margin = Mm(12.7)

    def generate_exam_bundle(
        self,
        *,
        teacher_id: Optional[int],
        class_id: int,
        subject: str,
        exam_type: str,
        num_questions: int,
        num_versions: int,
        level: str,
        student_id_columns: int = 8,
        duration_minutes: int = 60,
        semester: str = "Học kỳ II",
        academic_year: str = "Năm học 2025 - 2026",
        exam_date: str = "....... / ....... / ............",
        school_name: str = DEFAULT_SCHOOL_NAME,
        department_name: str = DEFAULT_DEPARTMENT_NAME,
        exam_title: str = DEFAULT_EXAM_TITLE,
    ) -> GeneratedBatchBundle:
        normalized_exam_type = _normalize_exam_type_key(exam_type or "")
        if normalized_exam_type != "trac_nghiem":
            raise ValueError("Luồng đề thi OCR hiện chỉ hỗ trợ đề trắc nghiệm A/B/C/D.")
        if num_questions < 1 or num_questions > 120:
            raise ValueError("Số câu hỏi phải nằm trong khoảng 1-120.")
        if num_versions < 1 or num_versions > 8:
            raise ValueError("Số mã đề phải nằm trong khoảng 1-8.")

        target_class = self.db.query(models.Classroom).filter_by(id=class_id).first()
        if not target_class:
            raise LookupError("Lớp học không tồn tại.")

        docs = (
            self.db.query(models.Document)
            .filter(
                models.Document.class_id == class_id,
                models.Document.subject_id == target_class.subject_id,
            )
            .all()
        )
        if not docs:
            raise LookupError("Lớp học này chưa có tài liệu phù hợp để sinh đề OCR.")

        allowed_filenames = [doc.filename for doc in docs if doc.filename]
        chunk_texts = _collect_chunk_texts(self.db, allowed_filenames, limit=240)
        context_summary = "\n".join(chunk_texts[:80])[:14000] if chunk_texts else f"Học liệu môn {subject}"

        question_pool = self._build_question_pool(
            subject=subject,
            exam_type=normalized_exam_type,
            num_questions=num_questions,
            target_class=target_class,
            docs=docs,
            context_summary=context_summary,
            chunk_texts=chunk_texts,
        )
        if len(question_pool) < num_questions:
            raise RuntimeError("Không thể chuẩn bị đủ số câu hỏi để xuất đề OCR.")

        ensure_test_ocr_dirs()
        batch_code = build_batch_code()
        file_path = build_generated_docx_path(batch_code)

        doc = Document()
        apply_document_style(doc)
        self._prepare_section(doc)

        layout = _build_omr_layout(question_count=num_questions, student_id_columns=student_id_columns, exam_code_columns=3)
        layout["sheet_meta"] = {
            "school_name": school_name,
            "department_name": department_name,
            "subject_name": subject,
            "class_name": target_class.name,
            "exam_date": exam_date,
        }

        omr_stream = _render_omr_sheet(
            version_code=None,
            question_count=num_questions,
            layout=layout,
            school_name=school_name,
            department_name=department_name,
            subject_name=subject,
            class_name=target_class.name,
            exam_date=exam_date,
        )
        picture_paragraph = doc.add_paragraph()
        picture_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        picture_paragraph.add_run().add_picture(omr_stream, width=Inches(7.7))
        picture_paragraph.space_after = Pt(0)

        used_codes: set[str] = set()
        answer_key_json: List[Dict[str, object]] = []
        answer_versions: List[dict] = []

        for version_index in range(num_versions):
            while True:
                exam_code = str(random.randint(101, 999))
                if exam_code not in used_codes:
                    used_codes.add(exam_code)
                    break

            selected_questions = copy.deepcopy(random.sample(question_pool, num_questions))
            random.shuffle(selected_questions)

            for question in selected_questions:
                options = copy.deepcopy(question.get("options", []))
                correct_idx = ord(question.get("ans", "A").upper()) - 65
                correct_text = options[correct_idx] if 0 <= correct_idx < len(options) else options[0]
                random.shuffle(options)
                question["options"] = options
                question["new_ans"] = ["A", "B", "C", "D"][options.index(correct_text)]

            add_school_header(
                doc,
                school_name=school_name,
                department_name=department_name,
                exam_title=exam_title,
                section_title="ĐỀ THI TRẮC NGHIỆM",
            )
            add_exam_info_table(
                doc,
                subject=subject,
                class_name=target_class.name,
                exam_code=exam_code,
                exam_type_label=DEFAULT_EXAM_TYPE_LABEL,
                semester=semester,
                academic_year=academic_year,
                exam_date=exam_date,
                duration_minutes=duration_minutes,
            )
            add_exam_instruction_block(doc, exam_type_label=DEFAULT_EXAM_TYPE_LABEL)
            add_multiple_choice_questions(doc, selected_questions)

            if version_index < num_versions - 1:
                doc.add_page_break()

            answer_list = [question.get("new_ans", question.get("ans", "A")) for question in selected_questions]
            answer_key_json.append(
                {
                    "version_index": version_index + 1,
                    "exam_code": exam_code,
                    "question_count": num_questions,
                    "answer_key": answer_list,
                    "questions": [question.get("q", "") for question in selected_questions],
                }
            )
            answer_versions.append({"code": exam_code, "questions": [{"new_ans": item} for item in answer_list]})

        add_answer_key_section(doc, answer_versions, exam_type_label=DEFAULT_EXAM_TYPE_LABEL)

        doc.save(str(file_path))
        answer_key_file_path = build_generated_answer_xlsx_path(batch_code)
        answer_key_file_path.write_bytes(build_answer_key_workbook(answer_key_json))

        batch = models.TestOCRExamBatch(
            teacher_id=teacher_id,
            class_id=class_id,
            subject_id=target_class.subject_id,
            subject_name=subject,
            exam_type="trac_nghiem",
            level=level,
            num_questions=num_questions,
            num_versions=num_versions,
            batch_code=batch_code,
            generated_docx_path=str(file_path),
            answer_key_json=answer_key_json,
            omr_layout_json=layout,
        )
        self.db.add(batch)
        self.db.commit()
        self.db.refresh(batch)

        return GeneratedBatchBundle(
            batch=batch,
            file_path=file_path,
            download_filename=f"DeThiOCR_{_safe_filename_fragment(subject)}_{num_versions}ma.docx",
            answer_key_file_path=answer_key_file_path,
            answer_key_download_filename=f"DapAnOCR_{_safe_filename_fragment(subject)}_{num_versions}ma.xlsx",
        )
