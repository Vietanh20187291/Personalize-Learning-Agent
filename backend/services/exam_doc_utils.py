from __future__ import annotations

import re
from typing import Iterable, Sequence

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

DEFAULT_SCHOOL_NAME = ""
DEFAULT_DEPARTMENT_NAME = ""
DEFAULT_EXAM_TITLE = "ĐỀ THI KẾT THÚC HỌC PHẦN"
DEFAULT_EXAM_TYPE_LABEL = "Trắc nghiệm"


def apply_document_style(doc: Document) -> None:
    section = doc.sections[-1]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _set_run_font(run, *, size: int = 12, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _clean_option_text(value: object) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    text = re.sub(r"^[\u2022\-\*]+\s*", "", text)
    text = re.sub(r"^(?:[A-Da-d])[\.\)\:\-]\s*", "", text)
    return text.strip()


def add_plain_paragraph(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    size: int = 12,
    align: WD_PARAGRAPH_ALIGNMENT = WD_PARAGRAPH_ALIGNMENT.LEFT,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=bold)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: int = 11,
    align: WD_PARAGRAPH_ALIGNMENT = WD_PARAGRAPH_ALIGNMENT.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_school_header(
    doc: Document,
    *,
    school_name: str,
    department_name: str,
    exam_title: str,
    section_title: str,
) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    left_cell = table.cell(0, 0)
    set_cell_text(
        left_cell,
        f"BỘ GIÁO DỤC VÀ ĐÀO TẠO\n{school_name}",
        bold=True,
        size=11,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
    )
    left_cell.add_paragraph()
    _set_run_font(left_cell.paragraphs[1].add_run(department_name), size=12, bold=True)
    left_cell.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    left_cell.add_paragraph()
    _set_run_font(left_cell.paragraphs[2].add_run("--------------"), size=12, bold=True)
    left_cell.paragraphs[2].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    right_cell = table.cell(0, 1)
    set_cell_text(
        right_cell,
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        bold=True,
        size=11,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
    )
    right_cell.add_paragraph()
    _set_run_font(right_cell.paragraphs[1].add_run("Độc lập - Tự do - Hạnh phúc"), size=12, bold=True)
    right_cell.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    right_cell.add_paragraph()
    _set_run_font(right_cell.paragraphs[2].add_run("--------------"), size=12, bold=True)
    right_cell.paragraphs[2].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()
    add_plain_paragraph(doc, exam_title, bold=True, size=16, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_plain_paragraph(doc, section_title, bold=True, size=13, align=WD_PARAGRAPH_ALIGNMENT.CENTER)


def add_exam_info_table(
    doc: Document,
    *,
    subject: str,
    class_name: str,
    exam_code: str,
    exam_type_label: str,
    semester: str,
    academic_year: str,
    exam_date: str,
    duration_minutes: int,
) -> None:
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows = [
        (f"Môn học: {subject}", f"Mã đề: {exam_code}"),
        (f"Lớp học phần: {class_name}", f"Hình thức thi: {exam_type_label}"),
        (semester, academic_year),
        (f"Ngày thi: {exam_date}", f"Thời gian làm bài: {duration_minutes} phút"),
    ]
    for row_index, (left_value, right_value) in enumerate(rows):
        set_cell_text(table.cell(row_index, 0), left_value, bold=row_index < 2)
        set_cell_text(table.cell(row_index, 1), right_value, bold=row_index < 2)

    doc.add_paragraph()


def add_candidate_block(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run(
        "Họ và tên sinh viên: ..............................................................    "
        "MSSV: ........................................\n"
        "Lớp: ..............................................................    "
        "Phòng thi: ...................................."
    )
    _set_run_font(run, size=12)


def add_exam_instruction_block(doc: Document, *, exam_type_label: str) -> None:
    add_plain_paragraph(doc, "Hướng dẫn làm bài:", bold=True)

    instructions = [
        "Sinh viên đọc kỹ đề thi trước khi làm bài.",
        "Không sử dụng tài liệu nếu giảng viên không cho phép.",
        "Chọn một đáp án đúng nhất cho mỗi câu hỏi.",
        "Ghi đáp án vào phiếu trả lời trắc nghiệm theo đúng mã đề được phát.",
    ]
    if exam_type_label.lower() != "trắc nghiệm":
        instructions = instructions[:2]
        instructions.append("Trình bày rõ ràng, mạch lạc và bám sát yêu cầu của từng câu hỏi.")

    for item in instructions:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        _set_run_font(run, size=12)
    doc.add_paragraph()


def add_multiple_choice_questions(doc: Document, questions: Sequence[dict]) -> None:
    for index, question in enumerate(questions, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        number_run = paragraph.add_run(f"Câu {index}. ")
        _set_run_font(number_run, size=12, bold=True)
        text_run = paragraph.add_run(question.get("q", ""))
        _set_run_font(text_run, size=12)

        for label, option in zip(["A", "B", "C", "D"], question.get("options", [])):
            option_paragraph = doc.add_paragraph()
            option_paragraph.paragraph_format.left_indent = Cm(0.8)
            option_paragraph.paragraph_format.space_after = Pt(1)
            option_run = option_paragraph.add_run(f"{label}. {_clean_option_text(option)}")
            _set_run_font(option_run, size=12)
        doc.add_paragraph()


def add_answer_sheet(
    doc: Document,
    *,
    school_name: str,
    department_name: str,
    exam_title: str,
    subject: str,
    class_name: str,
    exam_code: str,
    exam_type_label: str,
    semester: str,
    academic_year: str,
    exam_date: str,
    duration_minutes: int,
    num_questions: int,
) -> None:
    doc.add_page_break()
    add_school_header(
        doc,
        school_name=school_name,
        department_name=department_name,
        exam_title=exam_title,
        section_title="PHIẾU TRẢ LỜI TRẮC NGHIỆM",
    )
    add_exam_info_table(
        doc,
        subject=subject,
        class_name=class_name,
        exam_code=exam_code,
        exam_type_label=exam_type_label,
        semester=semester,
        academic_year=academic_year,
        exam_date=exam_date,
        duration_minutes=duration_minutes,
    )

    info_table = doc.add_table(rows=2, cols=2)
    info_table.style = "Table Grid"
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(info_table.cell(0, 0), "Họ và tên (viết chữ in hoa): ........................................................................")
    set_cell_text(info_table.cell(0, 1), "MSSV: ........................................................")
    set_cell_text(info_table.cell(1, 0), "Lớp: .............................................................................................................")
    set_cell_text(info_table.cell(1, 1), "Chữ ký sinh viên: ...........................................")

    note = doc.add_paragraph()
    _set_run_font(note.add_run("Lưu ý: "), bold=True)
    _set_run_font(
        note.add_run(
            "Sinh viên ghi duy nhất một chữ cái A, B, C hoặc D vào chỗ trống của từng câu hỏi. "
            "Viết rõ ràng, chữ in hoa, không dùng mực đỏ."
        )
    )
    doc.add_paragraph()

    total_rows = (num_questions + 3) // 4
    answer_table = doc.add_table(rows=total_rows + 1, cols=4)
    answer_table.style = "Table Grid"
    answer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col_index in range(4):
        set_cell_text(
            answer_table.cell(0, col_index),
            f"Cột trả lời {col_index + 1}",
            bold=True,
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        )

    question_number = 1
    for row_index in range(1, total_rows + 1):
        for col_index in range(4):
            if question_number <= num_questions:
                set_cell_text(answer_table.cell(row_index, col_index), f"Câu {question_number:02d}: ........................", size=12)
                question_number += 1
            else:
                set_cell_text(answer_table.cell(row_index, col_index), "")


def add_answer_key_section(doc: Document, versions: Iterable[dict], *, exam_type_label: str = DEFAULT_EXAM_TYPE_LABEL) -> None:
    doc.add_page_break()
    add_plain_paragraph(doc, "ĐÁP ÁN VÀ HƯỚNG DẪN CHẤM", bold=True, size=15, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    doc.add_paragraph()

    for version in versions:
        label = doc.add_paragraph()
        run = label.add_run(f"Mã đề {version['code']}")
        _set_run_font(run, bold=True)

        if exam_type_label.lower() == "trắc nghiệm":
            cols = 4
            questions = list(version["questions"])
            total_rows = (len(questions) + cols - 1) // cols
            table = doc.add_table(rows=total_rows, cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            question_index = 0
            for row_index in range(total_rows):
                for col_index in range(cols):
                    if question_index < len(questions):
                        question = questions[question_index]
                        set_cell_text(
                            table.cell(row_index, col_index),
                            f"Câu {question_index + 1:02d}: {question.get('new_ans', question.get('ans', 'A'))}",
                            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
                        )
                        question_index += 1
                    else:
                        set_cell_text(table.cell(row_index, col_index), "")

        doc.add_paragraph("----------------------------------------------------------------")
